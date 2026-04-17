"""DOCX → ParsedRFP adapter using python-docx.

DOCX gives us first-class heading + table access, so this adapter produces
richer output than the PDF path (no regex-based heading detection needed).
"""

from __future__ import annotations

import logging
from io import BytesIO
from pathlib import Path

from docx import Document

from parsers.models import ParsedRFP, Section, TableBlob

logger = logging.getLogger(__name__)


def _heading_level(style_name: str) -> int | None:
    """Return 1..6 for Word 'Heading 1'..'Heading 6', else None."""
    if not style_name:
        return None
    name = style_name.strip().lower()
    if not name.startswith("heading"):
        return None
    tail = name.replace("heading", "").strip()
    if tail.isdigit():
        level = int(tail)
        return level if 1 <= level <= 6 else None
    return None


def _table_to_text(table) -> str:  # type: ignore[no-untyped-def]
    rows: list[str] = []
    for row in table.rows:
        cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
        rows.append(" | ".join(cells))
    return "\n".join(rows)


def _metadata(doc) -> dict[str, str]:  # type: ignore[no-untyped-def]
    core = getattr(doc, "core_properties", None)
    if core is None:
        return {}
    return {
        key: str(getattr(core, key, "") or "")
        for key in ("title", "author", "subject", "keywords", "comments")
        if getattr(core, key, None)
    }


def parse_docx_bytes(data: bytes, source_filename: str) -> ParsedRFP:
    """Parse DOCX binary into ParsedRFP. Uses paragraph styles for real headings."""
    doc = Document(BytesIO(data))

    sections: list[Section] = []
    current = Section(heading="(preamble)", level=0, text="")
    raw_text_parts: list[str] = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.rstrip()
        raw_text_parts.append(text)
        level = _heading_level(paragraph.style.name if paragraph.style else "")
        if level is not None and text:
            if current.text.strip() or current.heading != "(preamble)":
                sections.append(current)
            current = Section(heading=text, level=level, text="")
        else:
            if text:
                current.text += text + "\n"
    if current.text.strip() or current.heading != "(preamble)":
        sections.append(current)

    tables = [TableBlob(raw_text=_table_to_text(t)) for t in doc.tables if t.rows]

    parsed = ParsedRFP(
        source_format="docx",
        source_filename=source_filename,
        page_count=None,  # DOCX has no fixed pagination
        sections=sections,
        tables=tables,
        raw_text="\n".join(raw_text_parts),
        metadata=_metadata(doc),
    )
    logger.info(
        "docx_adapter.done file=%s sections=%d tables=%d",
        source_filename,
        len(parsed.sections),
        len(parsed.tables),
    )
    return parsed


def parse_docx_path(path: str | Path, source_filename: str | None = None) -> ParsedRFP:
    p = Path(path)
    return parse_docx_bytes(p.read_bytes(), source_filename or p.name)


__all__ = ["parse_docx_bytes", "parse_docx_path"]
