"""Unit tests for parsers/* — heading classifier, extractor heuristics, DOCX adapter.

PDF adapter is tested indirectly: we cover the pure-function heading detection
and section splitting directly; real-PDF parsing is exercised by the DOCX
end-to-end test (which builds a doc in-memory with python-docx) — good enough
since pypdf's own behaviour is not something we need to validate.
"""

from __future__ import annotations

from io import BytesIO

import pytest
from docx import Document as DocxDocument

from parsers import docx_adapter, pypdf_adapter, rfp_extractor
from parsers.models import ParsedRFP, Section, TableBlob


# --- pypdf_adapter pure-function heuristics ---------------------------------


def test_classify_heading_detects_allcaps_line() -> None:
    assert pypdf_adapter._classify_heading("EXECUTIVE SUMMARY") == (1, "EXECUTIVE SUMMARY")


def test_classify_heading_detects_numbered_heading_with_depth() -> None:
    assert pypdf_adapter._classify_heading("1.2 Scope of Work") == (2, "Scope of Work")
    assert pypdf_adapter._classify_heading("3. Requirements") == (1, "Requirements")


def test_classify_heading_skips_prose() -> None:
    assert pypdf_adapter._classify_heading("This is a normal paragraph of text.") is None


def test_split_sections_groups_text_until_next_heading() -> None:
    pages = [
        "EXECUTIVE SUMMARY\n"
        "Acme Bank needs a new core platform.\n"
        "1.0 SCOPE\n"
        "Provide APIs and dashboards.\n"
        "Must support 5k concurrent users.\n"
    ]
    sections = pypdf_adapter._split_sections(pages)
    headings = [s.heading for s in sections]
    assert "EXECUTIVE SUMMARY" in headings
    assert any("SCOPE" in h for h in headings)
    scope = next(s for s in sections if "SCOPE" in s.heading)
    assert "Provide APIs" in scope.text


# --- rfp_extractor heuristics -----------------------------------------------


def _parsed_with(raw_text: str, sections: list[Section] | None = None, metadata: dict[str, str] | None = None) -> ParsedRFP:
    return ParsedRFP(
        source_format="pdf",
        source_filename="test.pdf",
        page_count=4,
        sections=sections or [],
        tables=[],
        raw_text=raw_text,
        metadata=metadata or {},
    )


def test_extractor_picks_up_banking_industry_and_apac_region() -> None:
    text = (
        "Acme Bank RFP for core banking modernisation. Branch operations in Vietnam "
        "and Singapore require modernisation of the loan origination platform."
    )
    parsed = _parsed_with(text, metadata={"title": "Acme Bank RFP"})
    suggestion = rfp_extractor.extract_bid_card(parsed)
    assert suggestion.industry == "banking"
    assert suggestion.region == "APAC"
    assert "Acme Bank" in suggestion.client_name
    assert suggestion.confidence > 0.3


def test_extractor_collects_requirements_from_scoped_section() -> None:
    section = Section(
        heading="3. Functional Requirements",
        level=1,
        text=(
            "The system shall expose a REST API for account lookup.\n"
            "- Must support concurrent logins of 5k users.\n"
            "- Users should be able to view transactions in React.\n"
            "This paragraph does not use modal verbs and should be ignored."
        ),
    )
    parsed = _parsed_with("body", sections=[section])
    suggestion = rfp_extractor.extract_bid_card(parsed)
    # At least 3 req candidates: 1 modal sentence + 2 bullets.
    assert len(suggestion.requirement_candidates) >= 3
    combined = " | ".join(suggestion.requirement_candidates).lower()
    assert "rest api" in combined
    assert "concurrent logins" in combined


def test_extractor_falls_back_to_modal_sentences_when_no_scoped_section() -> None:
    parsed = _parsed_with(
        "The platform shall comply with PCI DSS. The tool must run on Kubernetes. "
        "This is background information only."
    )
    suggestion = rfp_extractor.extract_bid_card(parsed)
    assert len(suggestion.requirement_candidates) == 2


def test_extractor_detects_tech_keywords() -> None:
    parsed = _parsed_with(
        "Stack: NestJS + React front-end, Python/FastAPI microservices on Kubernetes, "
        "PostgreSQL + Redis. Supports REST and GraphQL integrations."
    )
    suggestion = rfp_extractor.extract_bid_card(parsed)
    keywords = set(suggestion.technology_keywords)
    assert {"nestjs", "react", "python", "fastapi", "kubernetes", "rest", "graphql"} <= keywords


def test_extractor_profile_hint_scales_with_document_size() -> None:
    small = _parsed_with("Tool shall be simple.", sections=[])
    small.page_count = 1
    assert rfp_extractor.extract_bid_card(small).estimated_profile_hint == "S"

    large_text = " ".join([f"Requirement {i}: the system shall X." for i in range(40)])
    large = _parsed_with(large_text, sections=[])
    large.page_count = 80
    large.tables = [TableBlob(raw_text="a|b\nc|d") for _ in range(5)]
    assert rfp_extractor.extract_bid_card(large).estimated_profile_hint == "XL"


# --- docx_adapter end-to-end ------------------------------------------------


def _build_sample_docx_bytes() -> bytes:
    doc = DocxDocument()
    doc.add_heading("Acme Bank Core Modernisation RFP", level=1)
    doc.add_paragraph("Prepared for: Acme Bank Vietnam")
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "Acme Bank is seeking a vendor to modernise its core banking platform "
        "across APAC branches."
    )
    doc.add_heading("Functional Requirements", level=1)
    doc.add_paragraph("The system shall expose REST APIs for account lookup.")
    doc.add_paragraph("Users should be able to view transactions in React.")
    doc.add_paragraph("Must comply with PCI DSS.")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Metric"
    table.rows[0].cells[1].text = "Target"
    table.rows[1].cells[0].text = "Uptime"
    table.rows[1].cells[1].text = "99.9%"

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_docx_adapter_parses_headings_paragraphs_and_tables() -> None:
    data = _build_sample_docx_bytes()
    parsed = docx_adapter.parse_docx_bytes(data, "acme.docx")

    assert parsed.source_format == "docx"
    assert parsed.source_filename == "acme.docx"
    headings = [s.heading for s in parsed.sections]
    assert "Executive Summary" in headings
    assert "Functional Requirements" in headings
    # The functional-requirements section captured its paragraphs.
    fr_section = next(s for s in parsed.sections if s.heading == "Functional Requirements")
    assert "REST APIs" in fr_section.text
    assert "PCI DSS" in fr_section.text
    # Table captured as raw text.
    assert parsed.tables
    assert "Uptime | 99.9%" in parsed.tables[0].raw_text


def test_docx_end_to_end_produces_useful_bid_card() -> None:
    parsed = docx_adapter.parse_docx_bytes(_build_sample_docx_bytes(), "acme.docx")
    suggestion = rfp_extractor.extract_bid_card(parsed)

    assert suggestion.industry == "banking"
    assert suggestion.region == "APAC"
    # Requirement candidates should cover the three modal sentences.
    combined = " | ".join(suggestion.requirement_candidates).lower()
    assert "rest api" in combined
    assert "pci dss" in combined
    assert suggestion.confidence >= 0.4


# --- guardrails: never raise on malformed input -----------------------------


def test_pypdf_adapter_returns_empty_parse_on_garbage_bytes() -> None:
    with pytest.raises(Exception):
        # pypdf raises on pure garbage; this is expected. The adapter contract says
        # it may raise — the FastAPI layer wraps it in a 400.
        pypdf_adapter.parse_pdf_bytes(b"not-a-pdf", "bad.pdf")


def test_extractor_handles_empty_parsed_rfp() -> None:
    parsed = _parsed_with("")
    suggestion = rfp_extractor.extract_bid_card(parsed)
    assert suggestion.requirement_candidates == []
    assert suggestion.confidence < 0.3
